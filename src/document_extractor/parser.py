"""
Document Parser Module
Handles parsing of Word, TXT, and PDF files using PyMuPDF, python-docx, and pytesseract
"""

import os
from typing import Optional, Tuple
import fitz  # PyMuPDF
from docx import Document
import pytesseract
from PIL import Image
import io


class DocumentParser:
    """Parser for various document formats with OCR support"""
    
    def __init__(self, file_path: str, use_ocr: bool = False):
        """
        Initialize the parser with a file path
        
        Args:
            file_path: Path to the document file
            use_ocr: Whether to use OCR for scanned documents (default: False)
        """
        self.file_path = file_path
        self.file_extension = os.path.splitext(file_path)[1].lower()
        self.use_ocr = use_ocr
        
    def extract_text(self) -> Tuple[Optional[str], dict]:
        """
        Extract text from the document based on its format
        
        Returns:
            Tuple of (extracted text as string, metadata dict) or (None, {}) if parsing fails
        """
        try:
            if self.file_extension == '.docx':
                return self._extract_from_docx()
            elif self.file_extension == '.txt':
                return self._extract_from_txt()
            elif self.file_extension == '.pdf':
                return self._extract_from_pdf()
            else:
                raise ValueError(f"Unsupported file format: {self.file_extension}")
        except Exception as e:
            print(f"Error extracting text: {str(e)}")
            return None, {"error": str(e)}
    
    def _extract_from_docx(self) -> Tuple[str, dict]:
        """Extract text from Word document"""
        doc = Document(self.file_path)
        text_parts = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        text = '\n'.join(text_parts)
        metadata = {
            "format": "docx",
            "paragraphs": len([p for p in doc.paragraphs if p.text.strip()]),
            "tables": len(doc.tables),
        }
        
        return text, metadata
    
    def _extract_from_txt(self) -> Tuple[str, dict]:
        """Extract text from TXT file"""
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(self.file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                    metadata = {
                        "format": "txt",
                        "encoding": encoding,
                        "lines": len(text.split('\n')),
                    }
                    return text, metadata
            except UnicodeDecodeError:
                continue
        
        raise ValueError("Could not decode text file with any supported encoding")
    
    def _extract_from_pdf(self) -> Tuple[str, dict]:
        """Extract text from PDF file using PyMuPDF"""
        doc = fitz.open(self.file_path)
        text_parts = []
        total_pages = len(doc)
        pages_with_text = 0
        pages_with_images = 0
        
        for page_num in range(total_pages):
            page = doc[page_num]
            
            # Extract text
            page_text = page.get_text()
            
            if page_text.strip():
                text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
                pages_with_text += 1
            
            # Check for images and use OCR if needed
            if self.use_ocr:
                image_list = page.get_images()
                if image_list:
                    pages_with_images += 1
                    for img_index, img in enumerate(image_list):
                        try:
                            xref = img[0]
                            base_image = doc.extract_image(xref)
                            image_bytes = base_image["image"]
                            image = Image.open(io.BytesIO(image_bytes))
                            
                            # Perform OCR
                            ocr_text = pytesseract.image_to_string(image)
                            if ocr_text.strip():
                                text_parts.append(f"\n--- OCR Text from Page {page_num + 1}, Image {img_index + 1} ---\n{ocr_text}")
                        except Exception as e:
                            print(f"OCR error on page {page_num + 1}, image {img_index}: {str(e)}")
            
            # If no text found and OCR is enabled, try OCR on the whole page
            if not page_text.strip() and self.use_ocr:
                try:
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better OCR
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    ocr_text = pytesseract.image_to_string(img)
                    if ocr_text.strip():
                        text_parts.append(f"\n--- OCR Text from Page {page_num + 1} ---\n{ocr_text}")
                except Exception as e:
                    print(f"Page OCR error on page {page_num + 1}: {str(e)}")
        
        doc.close()
        
        text = '\n\n'.join(text_parts)
        metadata = {
            "format": "pdf",
            "total_pages": total_pages,
            "pages_with_text": pages_with_text,
            "pages_with_images": pages_with_images,
            "ocr_used": self.use_ocr,
        }
        
        return text, metadata
    
    def get_file_info(self) -> dict:
        """Get basic file information"""
        file_stats = os.stat(self.file_path)
        return {
            "file_name": os.path.basename(self.file_path),
            "file_size": file_stats.st_size,
            "file_extension": self.file_extension,
            "file_size_mb": round(file_stats.st_size / (1024 * 1024), 2),
        }

